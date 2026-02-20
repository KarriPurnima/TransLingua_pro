# import streamlit as st
# from langchain_core.prompts import ChatPromptTemplate
# from langchain_core.output_parsers import StrOutputParser
# from langchain_groq import ChatGroq
# from gtts import gTTS
# import base64
# from io import BytesIO
# import os
# from dotenv import load_dotenv
# import speech_recognition as sr
# from docx import Document
# # Load environment variables
# load_dotenv()
# groq_api_key = os.getenv("GROQ_API_KEY")  # Ensure your key is in .env

# # ------------------------
# # Helpers
# # ------------------------
# @st.cache_data
# def get_base64_of_bin_file(bin_file):
#     with open(bin_file, 'rb') as f:
#         return base64.b64encode(f.read()).decode()

# def set_png_as_page_bg(png_file):
#     bin_str = get_base64_of_bin_file(png_file)
#     st.markdown(f'''
#     <style>
#     body {{
#         background-image: url("data:image/png;base64,{bin_str}");
#         background-size: cover;
#     }}
#     </style>
#     ''', unsafe_allow_html=True)

# def translate_text(text):
#     groqApi = ChatGroq(model="llama-3.1-8b-instant", groq_api_key=groq_api_key, temperature=0)
#     outputparser = StrOutputParser()
#     chainSec = prompt | groqApi | outputparser
#     return chainSec.invoke({
#         'source_language': source_language,
#         'target_language': target_language,
#         'text': text
#     })

# def play_audio(text):
#     tts = gTTS(text=text, lang=language_codes[target_language], slow=False)
#     audio_buffer = BytesIO()
#     tts.write_to_fp(audio_buffer)
#     audio_buffer.seek(0)
#     audio_b64 = base64.b64encode(audio_buffer.read()).decode()
#     st.markdown(f'<audio controls><source src="data:audio/mp3;base64,{audio_b64}" type="audio/mp3"></audio>', unsafe_allow_html=True)

# # ------------------------
# # Background & Sidebar
# # ------------------------
# set_png_as_page_bg('img.jpg')
# st.sidebar.image("Translation made easy with TransLingua Pro.png", use_container_width=True)
# st.sidebar.title("Translator Settings")

# # Language selection
# language_codes = {
#     "English": "en", "Spanish": "es", "French": "fr", "German": "de",
#     "Chinese": "zh", "Japanese": "ja", "Korean": "ko", "Italian": "it",
#     "Portuguese": "pt", "Russian": "ru", "Arabic": "ar", "Hindi": "hi",
#     "Dutch": "nl", "Greek": "el", "Swedish": "sv", "Turkish": "tr",
#     "Vietnamese": "vi","Thai": "th",
#     "Indonesian": "id",
#     "Filipino": "fil",
#     "Swahili": "sw",
#     "Polish": "pl",
#     "Czech": "cs",
#     "Hungarian": "hu",
#     "Norwegian": "no",
#     "Finnish": "fi","Bengali": "bn",
#     "Urdu": "ur",
#     "Punjabi": "pa",
#     "Tamil": "ta",
#     "Telugu": "te",   
#     "Malayalam": "ml",
#     "Gujarati": "gu",
#     "Kannada": "kn",
#     "Marathi": "mr",
#     "Odia": "or"
# }
# languages = list(language_codes.keys())
# source_language = st.sidebar.selectbox("Source Language", languages)
# target_language = st.sidebar.selectbox("Target Language", languages)

# # Prompt template
# prompt = ChatPromptTemplate.from_messages(
#     [
#         ("system", "You are a language translator. Translate the following text from {source_language} to {target_language}."),
#         ("user", "Text: {text}")
#     ]
# )

# # ------------------------
# # Tabs for modes
# # ------------------------
# tab1, tab2, tab3 = st.tabs(["Text Translation", "File Translation", "Voice Translation"])

# # ------------------------
# # Tab 1: Text Translation
# # ------------------------
# with tab1:
#     st.header("Real-Time Text Translation")
#     input_text = st.text_area("Enter text to translate")
#     if st.button("Translate Text", key="text_translate") and input_text.strip():
#         translation = translate_text(input_text)
#         st.success(f"Translation ({source_language} → {target_language}):")
#         st.markdown(f"**{translation}**")
#         play_audio(translation)

# # ------------------------
# # Tab 2: File Upload Translation
# # ------------------------
# with tab2:
#     st.header("File Translation (.txt / .docx)")
#     uploaded_file = st.file_uploader("Upload a file", type=["txt", "docx"])
#     if uploaded_file:
#         if uploaded_file.name.endswith(".txt"):
#             content = uploaded_file.read().decode("utf-8")
#         else:
#             doc = Document(uploaded_file)
#             content = "\n".join([para.text for para in doc.paragraphs])

#         st.text_area("File Content", content, height=200)

#         if st.button("Translate File", key="file_translate") and content.strip():
#             translation = translate_text(content)
#             st.success(f"Translated content ({source_language} → {target_language}):")
#             st.markdown(f"**{translation}**")
#             play_audio(translation)

# # ------------------------
# # Tab 3: Voice Translation
# # ------------------------
# with tab3:
#     st.header("Voice Input Translation")
#     recognizer = sr.Recognizer()

#     if st.button("Start Recording", key="voice_record"):
#         with sr.Microphone() as source:
#             st.info("Recording... Speak now")
#             audio = recognizer.listen(source)
#         try:
#             input_text = recognizer.recognize_google(audio)
#             st.write("You said:", input_text)
#             translation = translate_text(input_text)
#             st.success(f"Translation ({source_language} → {target_language}):")
#             st.markdown(f"**{translation}**")
#             play_audio(translation)
#         except sr.UnknownValueError:
#             st.error("Could not understand audio")
#         except sr.RequestError as e:
#             st.error(f"Error: {e}")

import streamlit as st
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_groq import ChatGroq
from gtts import gTTS
import base64
from io import BytesIO
import os
from dotenv import load_dotenv
import speech_recognition as sr
from docx import Document

# Load environment variables
load_dotenv()

# Try .env first, then Streamlit secrets
groq_api_key = os.getenv("GROQ_API_KEY") or st.secrets.get("GROQ_API_KEY", None)

if not groq_api_key:
    st.error("⚠️ GROQ_API_KEY is not set. Please add it in Render's Environment Variables or Streamlit Secrets.")
    st.stop()

# ------------------------
# Helpers
# ------------------------
@st.cache_data
def get_base64_of_bin_file(bin_file):
    with open(bin_file, 'rb') as f:
        return base64.b64encode(f.read()).decode()

def set_png_as_page_bg(png_file):
    bin_str = get_base64_of_bin_file(png_file)
    st.markdown(f'''
    <style>
    body {{
        background-image: url("data:image/png;base64,{bin_str}");
        background-size: cover;
    }}
    </style>
    ''', unsafe_allow_html=True)

def translate_text(text):
    groqApi = ChatGroq(model="llama-3.1-8b-instant", groq_api_key=groq_api_key, temperature=0)
    outputparser = StrOutputParser()
    chainSec = prompt | groqApi | outputparser
    return chainSec.invoke({
        'source_language': source_language,
        'target_language': target_language,
        'text': text
    })

def play_audio(text):
    tts = gTTS(text=text, lang=language_codes[target_language], slow=False)
    audio_buffer = BytesIO()
    tts.write_to_fp(audio_buffer)
    audio_buffer.seek(0)
    audio_b64 = base64.b64encode(audio_buffer.read()).decode()
    st.markdown(f'<audio controls><source src="data:audio/mp3;base64,{audio_b64}" type="audio/mp3"></audio>', unsafe_allow_html=True)

# ------------------------
# Background & Sidebar
# ------------------------
try:
    set_png_as_page_bg('img.jpg')
except Exception:
    pass  # Skip background if image not found

try:
    st.sidebar.image("Translation made easy with TransLingua Pro.png", use_container_width=True)
except Exception:
    pass  # Skip sidebar image if not found

st.sidebar.title("Translator Settings")

# Language selection
language_codes = {
    "English": "en", "Spanish": "es", "French": "fr", "German": "de",
    "Chinese": "zh", "Japanese": "ja", "Korean": "ko", "Italian": "it",
    "Portuguese": "pt", "Russian": "ru", "Arabic": "ar", "Hindi": "hi",
    "Dutch": "nl", "Greek": "el", "Swedish": "sv", "Turkish": "tr",
    "Vietnamese": "vi", "Thai": "th", "Indonesian": "id", "Filipino": "fil",
    "Swahili": "sw", "Polish": "pl", "Czech": "cs", "Hungarian": "hu",
    "Norwegian": "no", "Finnish": "fi", "Bengali": "bn", "Urdu": "ur",
    "Punjabi": "pa", "Tamil": "ta", "Telugu": "te", "Malayalam": "ml",
    "Gujarati": "gu", "Kannada": "kn", "Marathi": "mr", "Odia": "or"
}
languages = list(language_codes.keys())
source_language = st.sidebar.selectbox("Source Language", languages)
target_language = st.sidebar.selectbox("Target Language", languages)

# Prompt template
prompt = ChatPromptTemplate.from_messages(
    [
        ("system", "You are a language translator. Translate the following text from {source_language} to {target_language}. Return only the translated text, nothing else."),
        ("user", "Text: {text}")
    ]
)

# ------------------------
# Tabs for modes
# ------------------------
tab1, tab2, tab3 = st.tabs(["Text Translation", "File Translation", "Voice Translation"])

# ------------------------
# Tab 1: Text Translation
# ------------------------
with tab1:
    st.header("Real-Time Text Translation")
    input_text = st.text_area("Enter text to translate")
    if st.button("Translate Text", key="text_translate") and input_text.strip():
        with st.spinner("Translating..."):
            try:
                translation = translate_text(input_text)
                st.success(f"Translation ({source_language} → {target_language}):")
                st.markdown(f"**{translation}**")
                play_audio(translation)
            except Exception as e:
                st.error(f"Translation failed: {e}")

# ------------------------
# Tab 2: File Upload Translation
# ------------------------
with tab2:
    st.header("File Translation (.txt / .docx)")
    uploaded_file = st.file_uploader("Upload a file", type=["txt", "docx"])
    if uploaded_file:
        if uploaded_file.name.endswith(".txt"):
            content = uploaded_file.read().decode("utf-8")
        else:
            doc = Document(uploaded_file)
            content = "\n".join([para.text for para in doc.paragraphs])

        st.text_area("File Content", content, height=200)

        if st.button("Translate File", key="file_translate") and content.strip():
            with st.spinner("Translating..."):
                try:
                    translation = translate_text(content)
                    st.success(f"Translated content ({source_language} → {target_language}):")
                    st.markdown(f"**{translation}**")
                    play_audio(translation)
                except Exception as e:
                    st.error(f"Translation failed: {e}")

# ------------------------
# Tab 3: Voice Translation
# ------------------------
with tab3:
    st.header("Voice Input Translation")
    recognizer = sr.Recognizer()

    if st.button("Start Recording", key="voice_record"):
        try:
            with sr.Microphone() as source:
                st.info("Recording... Speak now")
                audio = recognizer.listen(source, timeout=10)
            try:
                input_text = recognizer.recognize_google(audio)
                st.write("You said:", input_text)
                with st.spinner("Translating..."):
                    translation = translate_text(input_text)
                    st.success(f"Translation ({source_language} → {target_language}):")
                    st.markdown(f"**{translation}**")
                    play_audio(translation)
            except sr.UnknownValueError:
                st.error("Could not understand audio. Please try again.")
            except sr.RequestError as e:
                st.error(f"Speech recognition error: {e}")
        except Exception as e:
            st.error(f"Microphone error: {e}. Note: Voice input may not work on cloud deployments.")